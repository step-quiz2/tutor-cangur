"""
export.py — Stage 5 of the Cangur ingestion pipeline.

Regenerate the three human-readable export artefacts from the
authoritative `problems_<level>eso.py` source files:

    exports/classificacio.xlsx   — flat one-row-per-problem table
    exports/pistes_inicials.docx — initial hints per problem (no answers)
    exports/distractors.docx     — per-option mistake commentary

Usage:
    python export.py
    python export.py --output-dir exports

Run after Stage 4 completes for a whole level/year, or any time the
catalog changes. See PIPELINE.md, Stage 5 for the full spec.
"""

import argparse
import importlib
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side


LEVEL_MODULES: list[tuple[str, str]] = [
    ("1ESO", "problems_1eso"),
    ("2ESO", "problems_2eso"),
    ("3ESO", "problems_3eso"),
    ("4ESO", "problems_4eso"),
]

XLSX_HEADERS = ["id", "categoria", "any", "numero", "punts", "tema", "resposta_correcta"]


# ---------------------------------------------------------------------------
# Catalog loading
# ---------------------------------------------------------------------------

def load_all_problems() -> list[tuple[str, list[dict]]]:
    """Import each level module and return [(level, [problem, …]), …].

    Each problem list is sorted by (any, numero) for stable downstream output.
    """
    out: list[tuple[str, list[dict]]] = []
    for level, modname in LEVEL_MODULES:
        mod = importlib.import_module(modname)
        if not hasattr(mod, "PROBLEMS"):
            raise RuntimeError(f"{modname} has no PROBLEMS dict")
        problems = list(mod.PROBLEMS.values())
        problems.sort(key=lambda p: (p["any"], p["numero"]))
        out.append((level, problems))
    return out


# ---------------------------------------------------------------------------
# 5a — exports/classificacio.xlsx
# ---------------------------------------------------------------------------

def write_classificacio(per_level: list[tuple[str, list[dict]]], out_path: Path) -> int:
    """Write the flat classification spreadsheet. Returns number of data rows."""
    wb = Workbook()
    ws = wb.active
    ws.title = "classificacio"

    header_font = Font(name="Arial", bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", start_color="305496")
    header_align = Alignment(horizontal="center", vertical="center")
    body_font = Font(name="Arial")
    body_align_left = Alignment(horizontal="left", vertical="center")
    body_align_center = Alignment(horizontal="center", vertical="center")
    thin = Side(style="thin", color="BFBFBF")
    body_border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col_idx, header in enumerate(XLSX_HEADERS, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    # Stage 5a sort: categoria, any, numero
    all_rows: list[dict] = [p for _, probs in per_level for p in probs]
    all_rows.sort(key=lambda p: (p["categoria"], p["any"], p["numero"]))

    center_cols = {2, 3, 4, 5, 7}  # categoria, any, numero, punts, resposta_correcta
    for row_idx, p in enumerate(all_rows, start=2):
        values = [
            p["id"],
            p["categoria"],
            p["any"],
            p["numero"],
            p["punts"],
            p.get("tema"),
            p["resposta_correcta"],
        ]
        for col_idx, val in enumerate(values, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.font = body_font
            cell.border = body_border
            cell.alignment = body_align_center if col_idx in center_cols else body_align_left

    # Column widths (DXA-free since openpyxl uses character units)
    widths = {"A": 22, "B": 11, "C": 8, "D": 9, "E": 8, "F": 16, "G": 20}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    ws.row_dimensions[1].height = 22
    ws.freeze_panes = "A2"

    wb.save(out_path)
    return len(all_rows)


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _set_default_font(doc: Document) -> None:
    """Set Arial as the default font for the document."""
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)


def _add_title(doc: Document, text: str, subtitle: str) -> None:
    title = doc.add_heading(text, level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x36, 0x5D)
    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_after = Pt(10)


# ---------------------------------------------------------------------------
# 5b — exports/pistes_inicials.docx
# ---------------------------------------------------------------------------

def write_pistes(per_level: list[tuple[str, list[dict]]], out_path: Path) -> int:
    """Write the hints document. NEVER includes expected_reasoning or
    resposta_correcta — by design (PIPELINE.md Stage 5b).
    Returns number of entries written.
    """
    doc = Document()
    _set_default_font(doc)
    _add_title(
        doc,
        "Pistes inicials — Cangur",
        "Document per al facilitador del taller. Cada entrada conté la "
        "pista pre-escrita del catàleg. No conté ni la resposta correcta "
        "ni el raonament complet.",
    )

    n_entries = 0
    for level, problems in per_level:
        if not problems:
            continue
        doc.add_heading(level, level=1)
        for p in problems:
            head = doc.add_paragraph()
            head.paragraph_format.space_before = Pt(6)
            head.paragraph_format.space_after = Pt(0)
            id_run = head.add_run(p["id"])
            id_run.bold = True
            tema = p.get("tema") or "tema no especificat"
            head.add_run(f"  ({p['punts']} punts · {tema})")

            pista_p = doc.add_paragraph()
            pista_p.paragraph_format.left_indent = Pt(18)
            pista_p.paragraph_format.space_after = Pt(4)
            label = pista_p.add_run("Pista: ")
            label.italic = True
            pista_text = p.get("pista_inicial") or "(sense pista pre-escrita)"
            pista_p.add_run(pista_text)
            n_entries += 1

    doc.save(out_path)
    return n_entries


# ---------------------------------------------------------------------------
# 5c — exports/distractors.docx
# ---------------------------------------------------------------------------

def write_distractors(per_level: list[tuple[str, list[dict]]], out_path: Path) -> int:
    """Write the distractors document. The correct option is named at the
    top of each entry but never commented (PIPELINE.md Stage 5c).
    Returns number of entries written.
    """
    doc = Document()
    _set_default_font(doc)
    _add_title(
        doc,
        "Distractors estudiats — Cangur",
        "Document per al professorat. Per a cada problema, per què un "
        "alumne pot triar cadascuna de les quatre opcions incorrectes.",
    )

    n_entries = 0
    for level, problems in per_level:
        if not problems:
            continue
        doc.add_heading(level, level=1)
        for p in problems:
            head = doc.add_paragraph()
            head.paragraph_format.space_before = Pt(8)
            head.paragraph_format.space_after = Pt(2)
            head.add_run(p["id"]).bold = True
            head.add_run(f" — resposta correcta: {p['resposta_correcta']}")

            distractors = p.get("comentaris_distractors") or {}
            correct = p["resposta_correcta"]
            for letter in ("A", "B", "C", "D", "E"):
                if letter == correct:
                    continue
                entry = doc.add_paragraph(style="List Bullet")
                entry.paragraph_format.space_after = Pt(0)
                entry.add_run(f"{letter}: ").bold = True
                entry.add_run(distractors.get(letter, "(comentari no disponible)"))
            n_entries += 1

    doc.save(out_path)
    return n_entries


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------

def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[1])
    ap.add_argument(
        "--output-dir", type=Path, default=Path("exports"),
        help="Output directory (default: exports/)",
    )
    args = ap.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    per_level = load_all_problems()
    total = sum(len(probs) for _, probs in per_level)
    print(f"Loaded {total} problems across {len(per_level)} levels:")
    for level, probs in per_level:
        years = sorted({p["any"] for p in probs})
        print(f"  {level}: {len(probs)} problems, years={years}")

    xlsx_path = args.output_dir / "classificacio.xlsx"
    n_rows = write_classificacio(per_level, xlsx_path)
    print(f"\n[5a] {xlsx_path}: {n_rows} data rows")

    pistes_path = args.output_dir / "pistes_inicials.docx"
    n_pistes = write_pistes(per_level, pistes_path)
    print(f"[5b] {pistes_path}: {n_pistes} entries")

    distr_path = args.output_dir / "distractors.docx"
    n_distr = write_distractors(per_level, distr_path)
    print(f"[5c] {distr_path}: {n_distr} entries")

    # PIPELINE.md Stage 5 validation: counts must equal len(PROBLEMS) summed.
    errors = []
    if n_rows != total:
        errors.append(f"XLSX rows ({n_rows}) != total problems ({total})")
    if n_pistes != total:
        errors.append(f"pistes entries ({n_pistes}) != total problems ({total})")
    if n_distr != total:
        errors.append(f"distractors entries ({n_distr}) != total problems ({total})")
    if errors:
        for err in errors:
            print(f"  VALIDATION FAILED: {err}", file=sys.stderr)
        return 1

    print(
        f"\nValidation OK: {total} problems = {n_rows} XLSX rows "
        f"= {n_pistes} pistes entries = {n_distr} distractors entries."
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
